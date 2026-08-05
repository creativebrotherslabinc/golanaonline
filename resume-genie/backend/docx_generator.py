import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ACCENT = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x11, 0x18, 0x27)
MID = RGBColor(0x37, 0x41, 0x51)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def _add_border_bottom(paragraph, color_hex="D1D5DB"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    _add_border_bottom(p, "D1D5DB")
    return p


def generate_docx(content: dict) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)

    # ── Header ────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(content.get("candidate_name", "Resume"))
    nr.font.size = Pt(24)
    nr.font.bold = True
    nr.font.color.rgb = DARK

    role = content.get("target_role", "")
    if role:
        role_p = doc.add_paragraph()
        role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role_p.paragraph_format.space_after = Pt(6)
        rr = role_p.add_run(role)
        rr.font.size = Pt(11)
        rr.font.color.rgb = ACCENT

    # ── Professional Summary ───────────────────────────────
    if content.get("professional_summary"):
        _section_header(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(content["professional_summary"])
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = MID

    # ── Skills ─────────────────────────────────────────────
    skills = content.get("skills", [])
    if skills:
        _section_header(doc, "CORE COMPETENCIES")
        skills_text = "  •  ".join(skills)
        sp = doc.add_paragraph(skills_text)
        sp.paragraph_format.space_after = Pt(2)
        for run in sp.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = MID

    # ── Work Experience ────────────────────────────────────
    if content.get("work_experience"):
        _section_header(doc, "WORK EXPERIENCE")

        for exp in content["work_experience"]:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_after = Pt(0)
            title_r = title_p.add_run(exp.get("title", ""))
            title_r.font.size = Pt(10.5)
            title_r.font.bold = True
            title_r.font.color.rgb = DARK

            dates_r = title_p.add_run(f"   {exp.get('dates', '')}")
            dates_r.font.size = Pt(9.5)
            dates_r.font.color.rgb = GRAY
            dates_r.font.italic = True

            company_p = doc.add_paragraph(exp.get("company", ""))
            company_p.paragraph_format.space_after = Pt(3)
            for run in company_p.runs:
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = GRAY

            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                br = bp.add_run(bullet)
                br.font.size = Pt(10)
                br.font.color.rgb = MID

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Education ──────────────────────────────────────────
    if content.get("education"):
        _section_header(doc, "EDUCATION")
        for edu in content["education"]:
            ep = doc.add_paragraph()
            ep.paragraph_format.space_after = Pt(1)
            er = ep.add_run(edu.get("degree", ""))
            er.font.size = Pt(10.5)
            er.font.bold = True
            er.font.color.rgb = DARK

            inst_p = doc.add_paragraph(
                f"{edu.get('institution', '')}  |  {edu.get('year', '')}"
            )
            inst_p.paragraph_format.space_after = Pt(4)
            for run in inst_p.runs:
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = GRAY

    # ── Certifications ─────────────────────────────────────
    certs = content.get("certifications", [])
    if certs:
        _section_header(doc, "CERTIFICATIONS")
        for cert in certs:
            if cert.strip():
                cp = doc.add_paragraph(style="List Bullet")
                cp.paragraph_format.space_after = Pt(2)
                cr = cp.add_run(cert)
                cr.font.size = Pt(10)
                cr.font.color.rgb = MID

    # ── Cover Letter page ──────────────────────────────────
    doc.add_page_break()

    cl_title = doc.add_paragraph()
    cl_title.paragraph_format.space_after = Pt(4)
    cl_r = cl_title.add_run("COVER LETTER")
    cl_r.font.size = Pt(16)
    cl_r.font.bold = True
    cl_r.font.color.rgb = ACCENT
    _add_border_bottom(cl_title, "1A56DB")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    cover = content.get("cover_letter", "")
    for para in cover.split("\n\n"):
        if para.strip():
            cp = doc.add_paragraph(para.strip())
            cp.paragraph_format.space_after = Pt(8)
            for run in cp.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = MID

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
